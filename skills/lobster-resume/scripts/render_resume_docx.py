#!/usr/bin/env python3
"""Word DOCX renderer for tailored resumes.

Input JSON shape matches render_resume_pdf.py:
{
  "basics": {"name": "", "headline": "", "contact": []},
  "sidebar": [{"title": "", "items": [""]}],
  "summary": "",
  "sections": [
    {"title": "", "entries": [
      {"title": "", "date": "", "bullets": [""]}
    ]}
  ]
}
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ModuleNotFoundError as exc:
    raise SystemExit(
        "Missing DOCX dependency: python-docx. Install with `uv pip install python-docx` "
        "or run this script with the Codex bundled workspace Python."
    ) from exc


BLUE = RGBColor(37, 99, 235)
TEXT = RGBColor(17, 24, 39)
MUTED = RGBColor(71, 85, 105)
LIGHT = "EEF6FF"
LINE = "BFDBFE"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = LINE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_run(run, size: float, color: RGBColor = TEXT, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "DengXian")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_para(container, text: str = "", size: float = 9.0, color: RGBColor = TEXT, bold: bool = False, space_after: float = 1.5):
    para = container.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.05
    run = para.add_run(text)
    set_run(run, size, color, bold)
    return para


def add_section_heading(container, title: str) -> None:
    para = add_para(container, title, 10.5, BLUE, True, 2)
    para.paragraph_format.space_before = Pt(5)
    p_pr = para._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), LINE)
    border.append(bottom)
    p_pr.append(border)


def add_bullet(container, text: str) -> None:
    para = container.add_paragraph(style=None)
    para.style = "List Bullet"
    para.paragraph_format.left_indent = Cm(0.35)
    para.paragraph_format.first_line_indent = Cm(-0.18)
    para.paragraph_format.space_after = Pt(1.2)
    para.paragraph_format.line_spacing = 1.05
    run = para.add_run(text)
    set_run(run, 8.8, TEXT)


def add_sidebar(cell, data: dict[str, Any]) -> None:
    set_cell_shading(cell, "F8FAFC")
    basics = data.get("basics", {})
    name = add_para(cell, basics.get("name", ""), 18, BLUE, True, 1)
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if basics.get("headline"):
        add_para(cell, basics["headline"], 8.5, MUTED, False, 5)

    contact = basics.get("contact", [])
    if contact:
        add_section_heading(cell, "联系")
        for item in contact:
            add_para(cell, str(item), 7.8, TEXT, False, 1)

    for block in data.get("sidebar", []):
        add_section_heading(cell, block.get("title", ""))
        for item in block.get("items", []):
            text = str(item)
            if len(text) <= 14:
                table = cell.add_table(rows=1, cols=1)
                table.autofit = True
                skill_cell = table.cell(0, 0)
                set_cell_border(skill_cell)
                set_cell_shading(skill_cell, "FFFFFF")
                add_para(skill_cell, text, 7.5, TEXT, False, 0)
            else:
                add_para(cell, text, 7.5, TEXT, False, 1)


def add_main(cell, data: dict[str, Any]) -> None:
    basics = data.get("basics", {})
    title = data.get("title") or f"面向 {basics.get('headline', '目标岗位')} 的定制简历"
    add_para(cell, title, 14.5, TEXT, True, 5)
    if data.get("summary"):
        add_para(cell, data["summary"], 8.8, MUTED, False, 5)

    for section in data.get("sections", []):
        add_section_heading(cell, section.get("title", ""))
        for entry in section.get("entries", []):
            role = entry.get("title", "")
            date = entry.get("date", "")
            line = f"{role}    {date}" if date else role
            add_para(cell, line, 9.2, TEXT, True, 1.5)
            for bullet in entry.get("bullets", []):
                add_bullet(cell, str(bullet))


def render(data: dict[str, Any], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "DengXian")
    styles["Normal"].font.size = Pt(9)

    layout = doc.add_table(rows=1, cols=2)
    layout.autofit = False
    layout.columns[0].width = Cm(5.1)
    layout.columns[1].width = Cm(13.5)
    left = layout.cell(0, 0)
    right = layout.cell(0, 1)
    left.width = Cm(5.1)
    right.width = Cm(13.5)

    add_sidebar(left, data)
    add_main(right, data)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a tailored resume JSON to DOCX.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    data = json.loads(args.input.read_text(encoding="utf-8"))
    render(data, args.output)
    print(json.dumps({"output": str(args.output)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
